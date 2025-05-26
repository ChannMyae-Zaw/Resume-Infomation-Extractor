import pandas as pd
from docx import Document
from io import BytesIO

def parse_analysis_dataframe(data):
    analysis_df = pd.json_normalize(
        data,
        record_path=['education'],
        meta=['name', ['contact', 'phone'], ['contact', 'email']],
        errors='ignore'
    )
    analysis_df.rename(columns={
        'contact.phone': 'phone',
        'contact.email': 'email'
    }, inplace=True)
    cols = ['name'] + [col for col in analysis_df.columns if col != 'name']
    return analysis_df[cols]

def parse_summary_dataframe(data):
    education_bullets = "\n".join([
        f"- {edu['degree']} at {edu['institution']} (GPA: {edu.get('gpa', 'N/A')})"
        for edu in data["education"]
    ])
    return pd.DataFrame([{
        "Name": data["name"],
        "Phone": data["contact"]["phone"],
        "Email": data["contact"]["email"],
        "Education": education_bullets
    }])



def dataframe_to_docx(df, title="Data Export"):
    doc = Document()
    doc.add_heading(title, 0)

    # Add table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)

    # Save to BytesIO
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f
